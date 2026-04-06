"""
NEXUS AI Tutor - Document Processing Service
Handles PDF, Word, and text document processing with AI analysis
"""

import io
import asyncio
from pathlib import Path
from typing import Dict, List, Optional, Any, BinaryIO
from dataclasses import dataclass
import logging
from datetime import datetime

logger = logging.getLogger('nexus.document')


@dataclass
class DocumentChunk:
    """A chunk of processed document"""
    text: str
    page_number: int
    chunk_index: int
    embedding: Optional[List[float]] = None
    metadata: Optional[Dict] = None


@dataclass
class ProcessedDocument:
    """Result of document processing"""
    document_id: str
    filename: str
    file_type: str
    total_pages: int
    total_chunks: int
    chunks: List[DocumentChunk]
    summary: str
    key_concepts: List[str]
    extracted_images: List[Dict]
    extracted_tables: List[Dict]
    processing_time: float
    metadata: Dict


class DocumentProcessor:
    """
    Document Processing Service
    Handles PDF, Word, and text files with content extraction,
    summarization, and educational analysis
    """
    
    SUPPORTED_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'doc',
        'text/plain': 'txt',
        'text/markdown': 'md'
    }
    
    def __init__(self, config: Dict):
        self.config = config
        self.chunk_size = config.get('chunk_size', 1000)
        self.chunk_overlap = config.get('chunk_overlap', 200)
    
    async def process_document(
        self,
        file: BinaryIO,
        filename: str,
        content_type: str
    ) -> ProcessedDocument:
        """
        Process a document and extract educational content
        
        Args:
            file: File-like object
            filename: Original filename
            content_type: MIME type
        """
        start_time = datetime.now()
        
        # Determine file type
        file_type = self.SUPPORTED_TYPES.get(content_type)
        if not file_type:
            raise ValueError(f"Unsupported file type: {content_type}")
        
        # Extract raw content based on file type
        if file_type == 'pdf':
            content, images, tables, metadata = await self._process_pdf(file)
        elif file_type in ['docx', 'doc']:
            content, images, tables, metadata = await self._process_word(file)
        else:
            content, images, tables, metadata = await self._process_text(file)
        
        # Chunk the content
        chunks = self._chunk_content(content)
        
        # Generate embeddings for semantic search
        chunks_with_embeddings = await self._generate_embeddings(chunks)
        
        # Generate summary
        summary = await self._generate_summary(content)
        
        # Extract key concepts
        key_concepts = await self._extract_concepts(content)
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        return ProcessedDocument(
            document_id=self._generate_doc_id(),
            filename=filename,
            file_type=file_type,
            total_pages=metadata.get('pages', 1),
            total_chunks=len(chunks_with_embeddings),
            chunks=chunks_with_embeddings,
            summary=summary,
            key_concepts=key_concepts,
            extracted_images=images,
            extracted_tables=tables,
            processing_time=processing_time,
            metadata=metadata
        )
    
    async def _process_pdf(self, file: BinaryIO) -> tuple:
        """Extract content from PDF"""
        import pdfplumber
        
        content_parts = []
        images = []
        tables = []
        metadata = {'pages': 0}
        
        with pdfplumber.open(file) as pdf:
            metadata['pages'] = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages):
                # Extract text
                text = page.extract_text() or ""
                content_parts.append({
                    'page': page_num + 1,
                    'text': text
                })
                
                # Extract tables
                page_tables = page.extract_tables()
                for table in page_tables:
                    tables.append({
                        'page': page_num + 1,
                        'data': table
                    })
                
                # Extract images (simplified - would use page.images in production)
                if page.images:
                    for img_idx, img in enumerate(page.images):
                        images.append({
                            'page': page_num + 1,
                            'index': img_idx,
                            'bbox': img.get('bbox')
                        })
        
        full_content = "\n\n".join(
            f"[Page {p['page']}]\n{p['text']}" 
            for p in content_parts
        )
        
        return full_content, images, tables, metadata
    
    async def _process_word(self, file: BinaryIO) -> tuple:
        """Extract content from Word document"""
        from docx import Document
        
        doc = Document(file)
        
        content_parts = []
        images = []
        tables = []
        metadata = {'pages': 1}  # Word doesn't have clear page count
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                'index': table_idx,
                'data': table_data
            })
        
        full_content = "\n\n".join(content_parts)
        
        return full_content, images, tables, metadata
    
    async def _process_text(self, file: BinaryIO) -> tuple:
        """Extract content from plain text"""
        content = file.read()
        if isinstance(content, bytes):
            content = content.decode('utf-8')
        
        return content, [], [], {'pages': 1}
    
    def _chunk_content(self, content: str) -> List[DocumentChunk]:
        """Split content into overlapping chunks"""
        chunks = []
        
        # Simple sentence-aware chunking
        sentences = content.replace('\n', ' ').split('. ')
        
        current_chunk = []
        current_length = 0
        chunk_index = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            sentence_length = len(sentence)
            
            if current_length + sentence_length > self.chunk_size and current_chunk:
                # Save current chunk
                chunk_text = '. '.join(current_chunk) + '.'
                chunks.append(DocumentChunk(
                    text=chunk_text,
                    page_number=1,  # Would be calculated from content position
                    chunk_index=chunk_index
                ))
                chunk_index += 1
                
                # Keep overlap
                overlap_sentences = current_chunk[-2:] if len(current_chunk) > 2 else current_chunk
                current_chunk = overlap_sentences
                current_length = sum(len(s) for s in current_chunk)
            
            current_chunk.append(sentence)
            current_length += sentence_length
        
        # Don't forget the last chunk
        if current_chunk:
            chunk_text = '. '.join(current_chunk) + '.'
            chunks.append(DocumentChunk(
                text=chunk_text,
                page_number=1,
                chunk_index=chunk_index
            ))
        
        return chunks
    
    async def _generate_embeddings(
        self, 
        chunks: List[DocumentChunk]
    ) -> List[DocumentChunk]:
        """Generate embeddings for each chunk"""
        # Placeholder - will use sentence-transformers
        for chunk in chunks:
            chunk.embedding = [0.0] * 768  # Placeholder embedding
        return chunks
    
    async def _generate_summary(self, content: str) -> str:
        """Generate educational summary of document"""
        # Use LLM to summarize
        # Placeholder
        return "AI-generated summary of the document highlighting key educational points..."
    
    async def _extract_concepts(self, content: str) -> List[str]:
        """Extract key concepts from document"""
        # Use NLP/LLM to extract concepts
        # Placeholder
        return ['concept_1', 'concept_2', 'concept_3']
    
    def _generate_doc_id(self) -> str:
        """Generate unique document ID"""
        import uuid
        return str(uuid.uuid4())
    
    async def highlight_errors(
        self,
        document_content: str,
        topic: str
    ) -> Dict[str, Any]:
        """
        Analyze document for errors and misconceptions
        
        Args:
            document_content: The document text
            topic: The expected topic/subject
        """
        # Use LLM to identify errors
        # Placeholder
        return {
            'errors': [
                {
                    'type': 'factual',
                    'location': 'paragraph 3',
                    'original': 'The error text...',
                    'correction': 'The correct text...',
                    'explanation': 'Why this is incorrect...'
                }
            ],
            'misconceptions': [
                {
                    'concept': 'concept_name',
                    'issue': 'What the student misunderstands',
                    'clarification': 'Correct understanding'
                }
            ],
            'overall_accuracy': 0.85
        }
    
    async def generate_study_guide(
        self,
        document: ProcessedDocument,
        student_level: str = 'intermediate'
    ) -> Dict[str, Any]:
        """
        Generate a study guide from processed document
        """
        return {
            'title': f'Study Guide: {document.filename}',
            'summary': document.summary,
            'key_concepts': document.key_concepts,
            'learning_objectives': [
                'Understand concept A',
                'Apply concept B',
                'Analyze concept C'
            ],
            'study_sections': [
                {
                    'title': 'Section 1',
                    'content': 'Review notes...',
                    'questions': ['Question 1?', 'Question 2?']
                }
            ],
            'practice_problems': [],
            'additional_resources': []
        }
