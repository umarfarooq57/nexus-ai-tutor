"""
NEXUS AI Tutor - Report Generation Service
Generates downloadable PDF and Word reports with diagrams
"""

import asyncio
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from datetime import datetime
import io
import logging

logger = logging.getLogger('nexus.reports')


@dataclass
class ReportSection:
    """A section in the report"""
    title: str
    content: str
    charts: List[Dict]
    tables: List[Dict]
    images: List[bytes]
    level: int  # Heading level (1, 2, 3)


@dataclass
class GeneratedReport:
    """A generated report"""
    report_id: str
    title: str
    report_type: str
    format: str  # 'pdf' or 'docx'
    file_data: bytes
    generation_time: float
    metadata: Dict


class ReportGenerator:
    """
    Report Generation Service
    Creates professional PDF and Word reports with charts and diagrams
    """
    
    REPORT_TYPES = [
        'progress_report',
        'assessment_report',
        'weakness_analysis',
        'learning_summary',
        'course_certificate',
        'study_guide'
    ]
    
    def __init__(self, config: Dict):
        self.config = config
        self.templates_dir = config.get('templates_dir', 'templates/reports')
    
    async def generate_report(
        self,
        report_type: str,
        data: Dict,
        format: str = 'pdf',
        student_id: Optional[str] = None
    ) -> GeneratedReport:
        """
        Generate a report
        
        Args:
            report_type: Type of report to generate
            data: Data for the report
            format: Output format ('pdf' or 'docx')
            student_id: Optional student ID for personalization
        """
        start_time = datetime.now()
        
        if report_type not in self.REPORT_TYPES:
            raise ValueError(f"Unknown report type: {report_type}")
        
        # Build report sections based on type
        sections = await self._build_sections(report_type, data)
        
        # Generate the report in requested format
        if format == 'pdf':
            file_data = await self._generate_pdf(sections, report_type, data)
        elif format == 'docx':
            file_data = await self._generate_docx(sections, report_type, data)
        else:
            raise ValueError(f"Unsupported format: {format}")
        
        generation_time = (datetime.now() - start_time).total_seconds()
        
        return GeneratedReport(
            report_id=self._generate_id(),
            title=self._get_report_title(report_type, data),
            report_type=report_type,
            format=format,
            file_data=file_data,
            generation_time=generation_time,
            metadata={
                'generated_at': datetime.now().isoformat(),
                'student_id': student_id,
                'data_summary': str(data.keys())
            }
        )
    
    async def _build_sections(
        self, 
        report_type: str, 
        data: Dict
    ) -> List[ReportSection]:
        """Build report sections based on type"""
        
        if report_type == 'progress_report':
            return await self._build_progress_sections(data)
        elif report_type == 'assessment_report':
            return await self._build_assessment_sections(data)
        elif report_type == 'weakness_analysis':
            return await self._build_weakness_sections(data)
        elif report_type == 'learning_summary':
            return await self._build_summary_sections(data)
        elif report_type == 'study_guide':
            return await self._build_study_guide_sections(data)
        else:
            return []
    
    async def _build_progress_sections(self, data: Dict) -> List[ReportSection]:
        """Build sections for progress report"""
        sections = []
        
        # Executive Summary
        sections.append(ReportSection(
            title="Executive Summary",
            content=f"""
This progress report covers the learning journey from {data.get('start_date', 'N/A')} to {data.get('end_date', 'N/A')}.

**Key Highlights:**
- Overall Progress: {data.get('overall_progress', 0):.0%}
- Topics Completed: {data.get('topics_completed', 0)}
- Time Invested: {data.get('total_time', 0)} hours
- Current Streak: {data.get('streak', 0)} days
            """,
            charts=[
                {'type': 'line', 'title': 'Progress Over Time', 'data': data.get('progress_timeline', [])}
            ],
            tables=[],
            images=[],
            level=1
        ))
        
        # Topic Breakdown
        sections.append(ReportSection(
            title="Topic Mastery Breakdown",
            content="Detailed analysis of mastery levels across different topics:",
            charts=[
                {'type': 'radar', 'title': 'Skill Distribution', 'data': data.get('skill_radar', [])}
            ],
            tables=[
                {
                    'title': 'Topic Performance',
                    'headers': ['Topic', 'Mastery', 'Time Spent', 'Quizzes'],
                    'rows': data.get('topic_table', [])
                }
            ],
            images=[],
            level=1
        ))
        
        # Strengths and Weaknesses
        sections.append(ReportSection(
            title="Strengths and Areas for Improvement",
            content=f"""
**Strong Areas:**
{self._format_list(data.get('strengths', []))}

**Areas Needing Attention:**
{self._format_list(data.get('weaknesses', []))}
            """,
            charts=[],
            tables=[],
            images=[],
            level=1
        ))
        
        # Recommendations
        sections.append(ReportSection(
            title="AI-Powered Recommendations",
            content=f"""
Based on your learning patterns and performance, here are personalized recommendations:

{self._format_numbered_list(data.get('recommendations', []))}
            """,
            charts=[],
            tables=[],
            images=[],
            level=1
        ))
        
        return sections
    
    async def _build_assessment_sections(self, data: Dict) -> List[ReportSection]:
        """Build sections for assessment report"""
        sections = []
        
        # Quiz Summary
        sections.append(ReportSection(
            title="Assessment Summary",
            content=f"""
**Quiz: {data.get('quiz_title', 'Assessment')}**

- Score: {data.get('score', 0):.0%}
- Questions Answered: {data.get('answered', 0)} / {data.get('total', 0)}
- Time Taken: {data.get('time_taken', 0)} minutes
- Status: {"Passed ✅" if data.get('passed', False) else "Needs Review ⚠️"}
            """,
            charts=[
                {'type': 'donut', 'title': 'Answer Distribution', 'data': data.get('answer_distribution', [])}
            ],
            tables=[],
            images=[],
            level=1
        ))
        
        # Question Analysis
        sections.append(ReportSection(
            title="Question-by-Question Analysis",
            content="Detailed breakdown of your responses:",
            charts=[],
            tables=[
                {
                    'title': 'Questions',
                    'headers': ['#', 'Question', 'Your Answer', 'Correct?', 'Topic'],
                    'rows': data.get('question_details', [])
                }
            ],
            images=[],
            level=1
        ))
        
        return sections
    
    async def _build_weakness_sections(self, data: Dict) -> List[ReportSection]:
        """Build sections for weakness analysis report"""
        return [
            ReportSection(
                title="Weakness Analysis Report",
                content=f"""
This report identifies areas where focused practice can significantly improve your mastery.

**Error Pattern Summary:**
{self._format_list(data.get('error_patterns', []))}

**Recommended Focus Areas:**
{self._format_numbered_list(data.get('focus_areas', []))}
                """,
                charts=[],
                tables=[],
                images=[],
                level=1
            )
        ]
    
    async def _build_summary_sections(self, data: Dict) -> List[ReportSection]:
        """Build sections for learning summary"""
        return [
            ReportSection(
                title="Learning Summary",
                content=data.get('summary', 'No summary available.'),
                charts=[],
                tables=[],
                images=[],
                level=1
            )
        ]
    
    async def _build_study_guide_sections(self, data: Dict) -> List[ReportSection]:
        """Build sections for study guide"""
        sections = []
        
        for topic in data.get('topics', []):
            sections.append(ReportSection(
                title=topic.get('title', 'Topic'),
                content=topic.get('content', ''),
                charts=[],
                tables=[],
                images=[],
                level=2
            ))
        
        return sections
    
    async def _generate_pdf(
        self, 
        sections: List[ReportSection], 
        report_type: str,
        data: Dict
    ) -> bytes:
        """Generate PDF report using WeasyPrint"""
        from jinja2 import Template
        
        # Build HTML content
        html_content = self._build_html(sections, report_type, data)
        
        try:
            from weasyprint import HTML
            
            pdf_buffer = io.BytesIO()
            HTML(string=html_content).write_pdf(pdf_buffer)
            return pdf_buffer.getvalue()
            
        except ImportError:
            # Fallback if WeasyPrint not available
            logger.warning("WeasyPrint not available, returning placeholder PDF")
            return b'%PDF-1.4 placeholder'
    
    async def _generate_docx(
        self, 
        sections: List[ReportSection], 
        report_type: str,
        data: Dict
    ) -> bytes:
        """Generate Word document"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = Document()
            
            # Title
            title = doc.add_heading(self._get_report_title(report_type, data), 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Date
            date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            doc.add_page_break()
            
            # Sections
            for section in sections:
                doc.add_heading(section.title, section.level)
                
                # Add content
                for paragraph in section.content.strip().split('\n\n'):
                    doc.add_paragraph(paragraph.strip())
                
                # Add tables
                for table_data in section.tables:
                    if table_data.get('rows'):
                        table = doc.add_table(rows=1, cols=len(table_data['headers']))
                        table.style = 'Table Grid'
                        
                        # Header row
                        for i, header in enumerate(table_data['headers']):
                            table.rows[0].cells[i].text = header
                        
                        # Data rows
                        for row_data in table_data['rows']:
                            row = table.add_row()
                            for i, cell in enumerate(row_data):
                                row.cells[i].text = str(cell)
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
            
        except ImportError:
            logger.warning("python-docx not available, returning placeholder")
            return b'placeholder docx'
    
    def _build_html(
        self, 
        sections: List[ReportSection], 
        report_type: str,
        data: Dict
    ) -> str:
        """Build HTML for PDF generation"""
        title = self._get_report_title(report_type, data)
        
        sections_html = ""
        for section in sections:
            heading_tag = f"h{section.level + 1}"
            sections_html += f"""
            <{heading_tag}>{section.title}</{heading_tag}>
            <div class="section-content">
                {"".join(f"<p>{p}</p>" for p in section.content.strip().split("\n\n"))}
            </div>
            """
        
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{title}</title>
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                body {{
                    font-family: 'Segoe UI', Arial, sans-serif;
                    line-height: 1.6;
                    color: #333;
                }}
                h1 {{
                    color: #0c8de6;
                    border-bottom: 2px solid #0c8de6;
                    padding-bottom: 10px;
                }}
                h2 {{
                    color: #a855f7;
                    margin-top: 30px;
                }}
                .section-content {{
                    margin: 20px 0;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 10px;
                    text-align: left;
                }}
                th {{
                    background-color: #0c8de6;
                    color: white;
                }}
                .highlight {{
                    background-color: #f0f7ff;
                    padding: 15px;
                    border-radius: 5px;
                }}
            </style>
        </head>
        <body>
            <h1>{title}</h1>
            <p class="date">Generated: {datetime.now().strftime('%B %d, %Y')}</p>
            {sections_html}
        </body>
        </html>
        """
    
    def _get_report_title(self, report_type: str, data: Dict) -> str:
        """Get report title based on type"""
        titles = {
            'progress_report': 'Learning Progress Report',
            'assessment_report': f"Assessment Report: {data.get('quiz_title', 'Quiz')}",
            'weakness_analysis': 'Weakness Analysis Report',
            'learning_summary': 'Learning Summary',
            'course_certificate': f"Certificate of Completion: {data.get('course_title', 'Course')}",
            'study_guide': f"Study Guide: {data.get('topic', 'Topic')}"
        }
        return titles.get(report_type, 'Report')
    
    def _format_list(self, items: List[str]) -> str:
        """Format a list as bullet points"""
        return '\n'.join(f"- {item}" for item in items)
    
    def _format_numbered_list(self, items: List[str]) -> str:
        """Format a list as numbered items"""
        return '\n'.join(f"{i+1}. {item}" for i, item in enumerate(items))
    
    def _generate_id(self) -> str:
        """Generate unique ID"""
        import uuid
        return str(uuid.uuid4())
